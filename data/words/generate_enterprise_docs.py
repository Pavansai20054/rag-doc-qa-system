import os
import random
import string
from datetime import datetime, timedelta
from typing import Dict, List, Tuple

from docx import Document


OUTPUT_DIR_NAME = "enterprise_docs"
TOTAL_DOCS = 500
MIN_WORDS = 500
MAX_WORDS = 1200


DEPARTMENTS = [
	"Human Resources",
	"Information Technology",
	"Engineering",
	"Security",
	"Compliance",
	"Finance",
	"Procurement",
	"Customer Support",
	"Product",
	"Operations",
	"Legal",
	"Data Platform",
	"Cloud Infrastructure",
]

REGIONS = ["US", "EU", "APAC"]
POLICY_VERSIONS = ["1.0", "1.1", "1.2", "2.0", "2.1", "3.0"]
TECHNOLOGIES = [
	"Kubernetes",
	"PostgreSQL",
	"Kafka",
	"Terraform",
	"Azure AD",
	"Okta",
	"ServiceNow",
	"Datadog",
	"Prometheus",
	"Vault",
	"Redis",
	"Snowflake",
	"AWS",
	"GCP",
	"Azure",
]
EMPLOYEE_ROLES = [
	"Systems Engineer",
	"Security Analyst",
	"SRE",
	"Platform Engineer",
	"HR Business Partner",
	"Finance Analyst",
	"Compliance Manager",
	"Procurement Lead",
	"Support Specialist",
	"Product Manager",
	"Release Manager",
]
SYSTEMS = [
	"Identity Access Gateway",
	"Ticketing Platform",
	"Log Aggregation Cluster",
	"Payment Processing Service",
	"Customer Data Lake",
	"Monitoring Stack",
	"CI/CD Pipeline",
	"Endpoint Management",
	"CRM Platform",
	"ERP System",
	"Document Repository",
]
PLATFORMS = [
	"Confluence",
	"Jira",
	"GitHub Enterprise",
	"Slack",
	"Zoom",
	"Workday",
	"SAP",
	"Salesforce",
	"Google Workspace",
	"Microsoft 365",
]
FRAMEWORKS = [
	"SOX",
	"SOC 2",
	"ISO 27001",
	"GDPR",
	"HIPAA",
	"PCI DSS",
]


CATEGORIES: Dict[str, List[str]] = {
	"HR Policies": [
		"Purpose and Scope",
		"Eligibility",
		"Policy Requirements",
		"Process",
		"Exceptions and Escalation",
	],
	"IT Security": [
		"Overview",
		"Access Control",
		"Monitoring",
		"Incident Response",
		"Audit Requirements",
	],
	"Engineering Runbooks": [
		"Service Summary",
		"Dependencies",
		"Operational Procedures",
		"Rollback Strategy",
		"Escalation",
	],
	"Incident Reports": [
		"Incident Summary",
		"Timeline",
		"Impact Analysis",
		"Root Cause",
		"Follow-Up Actions",
	],
	"Employee Onboarding": [
		"Pre-Start Checklist",
		"Day One",
		"Access Provisioning",
		"Training",
		"Success Metrics",
	],
	"Compliance Documents": [
		"Compliance Objective",
		"Control Mapping",
		"Evidence Collection",
		"Testing Procedures",
		"Reporting",
	],
	"DevOps SOPs": [
		"Standard Operating Procedure",
		"Change Management",
		"Release Workflow",
		"Monitoring and Alerts",
		"Post-Release Review",
	],
	"API Documentation": [
		"Authentication",
		"Endpoints",
		"Request and Response",
		"Error Handling",
		"Rate Limits",
	],
	"Cloud Infrastructure Guides": [
		"Architecture Overview",
		"Networking",
		"Identity and Access",
		"Backup and DR",
		"Cost Optimization",
	],
	"Finance Policies": [
		"Policy Statement",
		"Approval Workflow",
		"Budget Controls",
		"Reporting",
		"Audit Trail",
	],
	"Procurement Policies": [
		"Procurement Scope",
		"Vendor Selection",
		"Contract Review",
		"Purchase Order",
		"Renewal and Termination",
	],
	"Customer Support Procedures": [
		"Intake",
		"Triage",
		"Resolution",
		"Escalation",
		"Customer Follow-Up",
	],
	"Internal Wiki Pages": [
		"Context",
		"How We Work",
		"Key Contacts",
		"Reference Links",
		"FAQ",
	],
	"Product Manuals": [
		"Product Overview",
		"Installation",
		"Configuration",
		"Troubleshooting",
		"Maintenance",
	],
	"Meeting Notes": [
		"Attendees",
		"Agenda",
		"Discussion",
		"Decisions",
		"Action Items",
	],
}


SENTENCE_TEMPLATES = [
	"This document aligns with {framework} requirements and supports {region} data residency expectations.",
	"Teams must log requests in {platform} and track outcomes in {system} for audit readiness.",
	"The {department} organization uses {technology} to enforce consistent service baselines.",
	"All changes require peer review and validation in the {system} before production promotion.",
	"Escalations should follow the on-call rotation with a response target of {sla} minutes.",
	"Evidence artifacts are retained in {platform} for at least {retention} months.",
	"The process integrates with {system} and depends on {technology} for telemetry signals.",
	"Regional teams in {region} follow the same policy with localized legal review.",
	"Stakeholders include the {role}, service owner, and compliance lead for approvals.",
	"Risk acceptance must be documented and revalidated each quarter.",
	"We prioritize customer impact, data integrity, and regulatory alignment.",
	"All incidents are tagged with severity and mapped to service tiers.",
	"Audit sampling occurs monthly and is summarized in leadership dashboards.",
	"The {department} team coordinates with {role} to maintain operational readiness.",
	"Systems are monitored via {technology} and alert routing in {platform}.",
]


COMMON_PARAGRAPHS = [
	"This internal document is intended for authorized employees and outlines standard guidance for enterprise operations.",
	"Where possible, teams should reuse existing patterns, templates, and approved tooling to reduce operational risk.",
	"Any deviation from the described process must be approved through the documented exception workflow.",
]


TABLE_HEADERS = ["Item", "Owner", "Status"]
TABLE_STATUSES = ["Planned", "In Progress", "Complete", "Blocked"]

SIZE_PROFILES = [
	{
		"para_sentences": (2, 4),
		"extra_para_chance": 0.65,
		"bullets": (3, 6),
		"table_chance": 0.25,
	},
	{
		"para_sentences": (1, 3),
		"extra_para_chance": 0.45,
		"bullets": (2, 4),
		"table_chance": 0.15,
	},
]


def slugify(value: str) -> str:
	allowed = string.ascii_letters + string.digits + "_"
	normalized = value.replace(" ", "_").replace("/", "_").replace("-", "_")
	cleaned = "".join(ch if ch in allowed else "_" for ch in normalized)
	while "__" in cleaned:
		cleaned = cleaned.replace("__", "_")
	return cleaned.strip("_")


def random_date() -> str:
	days_back = random.randint(1, 900)
	date_value = datetime.utcnow() - timedelta(days=days_back)
	return date_value.strftime("%Y-%m-%d")


def count_words(text: str) -> int:
	return len([word for word in text.split() if word.strip()])


def build_sentence(context: Dict[str, str]) -> str:
	template = random.choice(SENTENCE_TEMPLATES)
	return template.format(**context)


def build_paragraph(context: Dict[str, str], sentences: int) -> str:
	return " ".join(build_sentence(context) for _ in range(sentences))


def add_paragraph(doc: Document, text: str) -> int:
	doc.add_paragraph(text)
	return count_words(text)


def add_bullet_list(doc: Document, items: List[str]) -> int:
	total = 0
	for item in items:
		doc.add_paragraph(item, style="List Bullet")
		total += count_words(item)
	return total


def add_table(doc: Document, rows: List[Tuple[str, str, str]]) -> int:
	table = doc.add_table(rows=1, cols=3)
	header_cells = table.rows[0].cells
	for index, header in enumerate(TABLE_HEADERS):
		header_cells[index].text = header
	total = count_words(" ".join(TABLE_HEADERS))
	for row in rows:
		row_cells = table.add_row().cells
		for index, cell_text in enumerate(row):
			row_cells[index].text = cell_text
		total += count_words(" ".join(row))
	return total


def generate_context(category: str) -> Dict[str, str]:
	department = random.choice(DEPARTMENTS)
	region = random.choice(REGIONS)
	version = random.choice(POLICY_VERSIONS)
	technology = random.choice(TECHNOLOGIES)
	role = random.choice(EMPLOYEE_ROLES)
	system = random.choice(SYSTEMS)
	platform = random.choice(PLATFORMS)
	framework = random.choice(FRAMEWORKS)
	sla = str(random.choice([15, 30, 45, 60]))
	retention = str(random.choice([12, 18, 24, 36]))

	return {
		"category": category,
		"department": department,
		"region": region,
		"version": version,
		"technology": technology,
		"role": role,
		"system": system,
		"platform": platform,
		"framework": framework,
		"sla": sla,
		"retention": retention,
	}


def add_metadata(doc: Document, context: Dict[str, str], author: str) -> int:
	word_count = 0
	word_count += add_paragraph(doc, f"Department: {context['department']}")
	word_count += add_paragraph(doc, f"Author: {author}")
	word_count += add_paragraph(doc, f"Created date: {random_date()}")
	word_count += add_paragraph(doc, f"Region: {context['region']}")
	word_count += add_paragraph(doc, f"Version: {context['version']}")
	return word_count


def build_section_paragraphs(context: Dict[str, str], profile: Dict[str, float]) -> List[str]:
	paragraphs = []
	min_sentences, max_sentences = profile["para_sentences"]
	paragraphs.append(build_paragraph(context, sentences=random.randint(min_sentences, max_sentences)))
	if random.random() < profile["extra_para_chance"]:
		paragraphs.append(build_paragraph(context, sentences=random.randint(min_sentences, max_sentences)))
	return paragraphs


def build_bullets(context: Dict[str, str], profile: Dict[str, float]) -> List[str]:
	bullets = []
	min_bullets, max_bullets = profile["bullets"]
	for _ in range(random.randint(min_bullets, max_bullets)):
		bullets.append(build_sentence(context))
	return bullets


def build_table_rows(context: Dict[str, str]) -> List[Tuple[str, str, str]]:
	rows = []
	for _ in range(random.randint(2, 4)):
		item = random.choice([
			"Access review",
			"Backup validation",
			"Patch window",
			"Vendor check",
			"Incident drill",
		])
		owner = random.choice([
			context["department"],
			context["role"],
			"Security",
			"Operations",
		])
		status = random.choice(TABLE_STATUSES)
		rows.append((item, owner, status))
	return rows


def generate_document(doc_index: int, output_dir: str) -> str:
	file_name = None
	output_path = None

	for attempt in range(6):
		profile = SIZE_PROFILES[min(attempt, len(SIZE_PROFILES) - 1)]
		category = random.choice(list(CATEGORIES.keys()))
		context = generate_context(category)
		author = (
			f"{random.choice(['Alex', 'Jordan', 'Casey', 'Morgan', 'Taylor', 'Avery', 'Riley'])} "
			f"{random.choice(['Nguyen', 'Patel', 'Kim', 'Garcia', 'Brown', 'Smith', 'Lee'])}"
		)
		title = f"{category} - {context['department']} {context['region']} v{context['version']}"
		file_base = slugify(f"{category}_{context['department']}_{context['region']}")
		file_name = f"{file_base}_{doc_index:03d}.docx"

		doc = Document()
		doc.add_heading(title, level=0)

		words_total = 0
		words_total += add_metadata(doc, context, author)

		for paragraph in COMMON_PARAGRAPHS:
			words_total += add_paragraph(doc, paragraph)

		sections = CATEGORIES[category]
		if attempt >= 3:
			max_sections = min(4, len(sections))
			sections = random.sample(sections, k=random.randint(3, max_sections))

		for section in sections:
			doc.add_heading(section, level=1)
			for paragraph in build_section_paragraphs(context, profile):
				words_total += add_paragraph(doc, paragraph)
			words_total += add_bullet_list(doc, build_bullets(context, profile))
			if random.random() < profile["table_chance"]:
				words_total += add_table(doc, build_table_rows(context))

		if words_total < MIN_WORDS:
			safety = 0
			while words_total < MIN_WORDS and safety < 6:
				doc.add_heading("Additional Notes", level=1)
				words_total += add_paragraph(doc, build_paragraph(context, sentences=3))
				words_total += add_bullet_list(doc, build_bullets(context, profile))
				safety += 1

		if MIN_WORDS <= words_total <= MAX_WORDS:
			output_path = os.path.join(output_dir, file_name)
			doc.save(output_path)
			return output_path

	if output_path is None:
		output_path = os.path.join(output_dir, file_name or f"doc_{doc_index:03d}.docx")
		doc.save(output_path)
	return output_path


def main() -> None:
	script_dir = os.path.dirname(os.path.abspath(__file__))
	output_dir = os.path.join(script_dir, OUTPUT_DIR_NAME)
	os.makedirs(output_dir, exist_ok=True)

	random.seed()
	for index in range(1, TOTAL_DOCS + 1):
		output_path = generate_document(index, output_dir)
		print(f"[{index}/{TOTAL_DOCS}] Created {os.path.basename(output_path)}")


if __name__ == "__main__":
	main()
